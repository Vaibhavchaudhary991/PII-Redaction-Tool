"""
main.py — Redact a .docx file.

Reads a Word document, replaces all required PII with fake data (Faker),
preserves bold/italic formatting, and saves a redacted copy.
Handles paragraphs AND tables.

Usage:  python main.py input.docx output.docx
"""

import sys
from docx import Document

from redactor import (
    REQUIRED_ENTITIES,
    FAKE_VALUES,
    create_analyzer,
    dedupe_overlaps,
    filter_results,
)


def redact_paragraph(paragraph, analyzer):
    """Detect PII in a paragraph and replace it, preserving run-level formatting."""
    text = paragraph.text
    if not text.strip():
        return

    results = analyzer.analyze(text=text, language="en")
    results = filter_results(results, text)
    results = dedupe_overlaps(results)
    if not results:
        return

    # One fake replacement per detected entity
    replacements = {(r.start, r.end): FAKE_VALUES[r.entity_type]() for r in results}

    # Apply at Run level so bold/italic styling is kept
    offset = 0
    replaced = set()
    for run in paragraph.runs:
        orig_len = len(run.text)
        run_start, run_end = offset, offset + orig_len
        new_run_text = run.text

        for r in sorted(results, key=lambda x: x.start, reverse=True):
            if r.end <= run_start or r.start >= run_end:
                continue
            local_start = max(r.start, run_start) - run_start
            local_end = min(r.end, run_end) - run_start
            key = (r.start, r.end)
            if key in replaced:
                # entity already replaced in an earlier run -> just remove overlap
                new_run_text = new_run_text[:local_start] + new_run_text[local_end:]
            else:
                new_run_text = (new_run_text[:local_start]
                                + replacements[key]
                                + new_run_text[local_end:])
                replaced.add(key)

        run.text = new_run_text
        offset += orig_len


def redact_document(doc, analyzer):
    """Redact paragraphs and table cells."""
    for para in doc.paragraphs:
        redact_paragraph(para, analyzer)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    redact_paragraph(para, analyzer)


def main():
    if len(sys.argv) != 3:
        print("Usage: python main.py input.docx output.docx")
        sys.exit(1)

    input_file, output_file = sys.argv[1], sys.argv[2]
    analyzer = create_analyzer()

    doc = Document(input_file)
    redact_document(doc, analyzer)
    doc.save(output_file)
    print(f"Redacted document saved to: {output_file}")


if __name__ == "__main__":
    main()
